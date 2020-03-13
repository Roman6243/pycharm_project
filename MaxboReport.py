import pandas as pd
from dfcleaner import cleaner
from docx import Document
from docx.oxml.shared import OxmlElement, qn
from docx.shared import Pt
from docx.shared import Inches
import numpy as np
import requests
import os
os.chdir(r'C:\Users\PC1\Dropbox (BigBlue&Company)\ETC Insight\Projects\Maxbo Report')
pd.set_option('display.max_columns', 30)
pd.set_option('display.width', 1000)

last_week_dates = pd.date_range(start='2020-03-02', end='2020-03-07', freq='D').strftime(date_format='%Y-%m-%d')
data = pd.DataFrame()
for date in last_week_dates:
    activity=requests.get("https://maxbo.link.express/external/api/v2/5d02982d29512bcc1729bb3964efb830/activity/query/?activity_date="+date+"T00:00:00&store_alias=ALL").json()
    for i in range(len(activity['store'])):
        store_name=activity['store'][i]['store_name']
        region_name=activity['store'][i]['region_name']
        user_count=activity['store'][i]['userCount']
        active_user_count=activity['store'][i]['activeUserCount']
        active_user_count_today=activity['store'][i]['activeTodayUserCount']
        temp=pd.DataFrame(activity['store'][i]['users']).assign(store_name=store_name, region_name=region_name, user_count=user_count, active_user_count=active_user_count,
                                                                active_user_count_today=active_user_count_today, date=date)
        data=pd.concat([data, temp])
    print('Extract data from %s'%date)
data.columns=cleaner.sanitize(data.columns) # cleaning the data headers

stores_using_system=data[data['active_today']][['store_name', 'region_name']].drop_duplicates().assign(is_active='yes') # all stores which used system at least once
stores_not_using_system=data[~data['store_name'].isin(stores_using_system['store_name'])][['store_name', 'region_name']].drop_duplicates().assign(is_active='no') # all stores which not used system even once
activities=pd.concat([stores_using_system, stores_not_using_system], axis=0).reset_index(drop=True)

print('Average use of the system (per user/per store/per region/per week): ')
n_usage_per_user_store_day=data[data['session_count']>0].groupby(['region_name', 'store_name', 'date'])['session_count'].mean().reset_index().\
                            groupby('region_name').mean().reset_index().rename(columns={'session_count':'n_user_store_day'})
n_not_active_stores=stores_not_using_system.groupby('region_name')['store_name'].count().reset_index().sort_values(by='region_name').rename(columns={'store_name':'n_not_active_stores'})
data_table_2=pd.merge(n_usage_per_user_store_day, n_not_active_stores, on='region_name')
data_table_2['score_goal']=data_table_2['n_user_store_day']-2
data_table_2['Goal'] = 2
data_table_2.loc['mean']=data_table_2.mean()
data_table_2.reset_index(drop=True, inplace=True)
data_table_2.loc[data_table_2['region_name'].isnull(), 'region_name'] = 'Average'
data_table_2=data_table_2.pivot_table(columns='region_name')
data_table_2=data_table_2.reset_index()
data_table_2.loc[data_table_2['index']=='n_not_active_stores', 'index'] = 'Number of stores that have not been using the system'
data_table_2.loc[data_table_2['index']=='n_user_store_day', 'index'] = 'Average use per day (nr of times)'
data_table_2.loc[data_table_2['index']=='score_goal', 'index'] = 'Score against goal'
data_table_2.rename(columns={'index': 'Region'}, inplace=True)
data_table_2['Region'] = pd.Categorical(data_table_2['Region'], ["Average use per day (nr of times)", 'Goal', "Score against goal", "Number of stores that have not been using the system"])
data_table_2.sort_values('Region', inplace=True)

# make a shading for particular cells in table
def shade_cells(cells, shade):
    for cell in cells:
        tcPr = cell._tc.get_or_add_tcPr()
        tcVAlign = OxmlElement("w:shd")
        tcVAlign.set(qn("w:fill"), shade)
        tcPr.append(tcVAlign)

dateparse = lambda x: pd.datetime.strptime(x, '%Y-%m-%d %H:%M:%S')
excluded_counter = pd.read_csv('C:/Users/PC1/Dropbox (BigBlue&Company)/ETC Insight/Projects/Maxbo_4 (python script)/input_files/excluded_counters.csv')
store_rename = pd.read_csv('C:/Users/PC1/Dropbox (BigBlue&Company)/ETC Insight/Projects/Maxbo_4 (python script)/input_files/store renaming.csv')
calendar_weeks = pd.read_excel('C:/Users/PC1/Dropbox (BigBlue&Company)/ETC Insight/Projects/Maxbo_4 (python script)/input_files/calendar weeks.xlsx', parse_dates=['Monday', 'Tuesday', 'Wednesday', 'Thursday', 'Friday', 'Saturday', 'Sunday'], date_parser=dateparse)
calendar_weeks = calendar_weeks.melt(id_vars = ['Calendar week', 'Year'], value_vars = ['Monday', 'Tuesday', 'Wednesday', 'Thursday', 'Friday', 'Saturday', 'Sunday'], var_name = 'Day', value_name='Date')
calendar_weeks = calendar_weeks.rename(columns = {'Calendar week': 'calendar_week', 'Year': 'calendar_year', 'Date':'date', 'Day':'calendar_day'})
store_info = pd.read_excel('C:/Users/PC1/Dropbox (BigBlue&Company)/ETC Insight/Projects/Maxbo_4 (python script)/input_files/stores_info.xlsx')
sales_counters = store_info[['store id', 'store_name', 'entrance name', 'type_counter', 'nr_counter', 'ip', 'counters_used', 'type_sales']]
sales_group = store_info[['store_name', 'type_sales', 'cr_group']].drop_duplicates()

# working hours formatting start
working_hours = store_info[['store_name', 'MT_From', 'MT_To', 'F_From', 'F_To', 'S_From', 'S_To']].drop_duplicates()
working_hours = working_hours.melt(id_vars='store_name', value_vars=['MT_From', 'MT_To', 'F_From', 'F_To', 'S_From', 'S_To'], value_name='hour', var_name='day_time')
working_hours = working_hours.fillna(0)
working_hours[['hour','min']]  = working_hours['hour'].astype(str).str.split(':', expand = True)
working_hours['min'] = working_hours['min'].fillna('0').astype(int)
working_hours['hour'] = working_hours['hour'].astype(int) + 1
working_hours = working_hours.drop(columns='min')
working_hours = pd.pivot_table(working_hours, index='store_name', columns='day_time').reset_index()
working_hours.columns = working_hours.columns.droplevel()
working_hours.columns = ['store_name', 'F_From', 'F_To', 'MT_From', 'MT_To', 'S_From', 'S_To']
# working hours formatting end

start = '2020-01-06'
end = '2020-03-07'

date_start=pd.date_range(start=start, end=end, freq='W-MON').strftime(date_format='%Y-%m-%d')
date_end=pd.date_range(start=start, end=end, freq='W-SAT').strftime(date_format='%Y-%m-%d')

data_sales_all = pd.DataFrame()
data_traffic_all = pd.DataFrame()
for next_date in range(len(date_start)):
    sales_req = requests.get("https://maxbo.link.express/external/api/v2/5d02982d29512bcc1729bb3964efb830/receipts/query/?store_alias=all&start_date=" +
                        date_start[next_date] + "T06:00:00&end_date=" + date_end[next_date] + "T23:00:00").json()
    data_sales  = pd.DataFrame()
    for j in range(len(sales_req['store'])):
        store_name = sales_req['store'][j]['store_name']
        for k in range(len(sales_req['store'][j]['dayStat'])):
            if sales_req['store'][j]['dayStat'][k] != None:
                date = sales_req['store'][j]['dayStat'][k]['date']
                temp_pos = pd.DataFrame(sales_req['store'][j]['dayStat'][k]['pos']['hourlySales']).assign(type_sales = 'pos', date = date, store_name = store_name)
                temp_sales = pd.DataFrame(sales_req['store'][j]['dayStat'][k]['sales']['hourlySales']).assign(type_sales = 'sales', date = date, store_name = store_name)
                data_sales = pd.concat([data_sales, temp_pos, temp_sales])
    data_sales['store_name'] = data_sales['store_name'].str.upper()
    data_sales['date'] = pd.to_datetime(data_sales['date'], format = '%Y-%m-%d').dt.normalize()
    data_sales_all = pd.concat([data_sales_all, data_sales])

    ### traffic data
    traffic_req = requests.get('https://maxbo.retailflux.com/traffic/api/v2/json/487A27B9BC3E29038F7E813365A417B1/query/?store_alias=all&start_date=' +
                         date_start[next_date] + "T06:00:00&end_date="+ date_end[next_date] + "T23:00:00").json()
    data_traffic = pd.DataFrame()
    for j in range(len(traffic_req['command_output']['store_camera_list'])):
        store_name = traffic_req['command_output']['store_camera_list'][j]['name']
        store_tags = traffic_req['command_output']['store_camera_list'][j]['tags']
        for k in range(len(traffic_req['command_output']['store_camera_list'][j]['counter_list'])):
            if traffic_req['command_output']['store_camera_list'][j]['counter_list'][k] != None:
                camera_name = traffic_req['command_output']['store_camera_list'][j]['counter_list'][k]['camera_name']
                counter_type = traffic_req['command_output']['store_camera_list'][j]['counter_list'][k]['counter_type']
                temp = pd.DataFrame(traffic_req['command_output']['store_camera_list'][j]['counter_list'][k]['data']).assign(store_name = store_name, counter = camera_name, store_tags = store_tags, counter_type = counter_type)
                data_traffic = pd.concat([data_traffic, temp])
    data_traffic['store_name'] = data_traffic['store_name'].str.upper()
    data_traffic['counter'] = data_traffic['counter'].replace(" ", "")
    data_traffic = data_traffic[~data_traffic['counter'].isin(excluded_counter['Counter'])]
    data_traffic['datetime'] = pd.to_datetime(data_traffic['datetime'], format = '%Y-%m-%dT%H:%M')
    data_traffic[['store_id', 'type_counter', 'nr_counter', 'ip']] = data_traffic['counter'].astype(str).str.split('_', expand = True)
    # store renaming
    for r in range(len(store_rename)):  data_traffic.loc[data_traffic['store_name'] == store_rename['It is'][r], 'store_name'] = store_rename['Should be'][r]
    data_traffic['hour'] = data_traffic['datetime'].dt.hour
    data_traffic['weekday'] = data_traffic['datetime'].dt.weekday
    data_traffic['year'] = data_traffic['datetime'].dt.year
    data_traffic['week'] = data_traffic['datetime'].dt.weekofyear
    data_traffic['date'] = data_traffic['datetime'].dt.normalize()
    data_traffic['weekdayname'] = data_traffic['datetime'].dt.weekday_name
    data_traffic = pd.merge(data_traffic, working_hours, on = 'store_name')
    # # select only work hours
    data_traffic.loc[(data_traffic['hour'] >= data_traffic['MT_From']) & (data_traffic['hour'] <= data_traffic['MT_To']) &
                     (data_traffic['weekday'] >= 0) & (data_traffic['weekday'] <= 3), 'open'] = True
    data_traffic.loc[(data_traffic['hour'] >= data_traffic['F_From']) & (data_traffic['hour'] <= data_traffic['F_To']) &
                     (data_traffic['weekday'] == 4), 'open'] = True
    data_traffic.loc[(data_traffic['hour'] >= data_traffic['S_From']) & (data_traffic['hour'] <= data_traffic['S_To']) &
                     (data_traffic['weekday'] == 5), 'open'] = True
    data_traffic = data_traffic[data_traffic['open'] == True] # choose only working hours
    data_traffic = pd.merge(data_traffic, sales_counters, on = ['ip', 'store_name', 'type_counter'], how='left')
    data_traffic = pd.merge(data_traffic, calendar_weeks, on = 'date', how = 'left')
    data_traffic_all = pd.concat([data_traffic_all, data_traffic])
    print('The start point is %s'%date_start[next_date])

data_traffic =  data_traffic_all[['people_in','store_name', 'type_counter', 'hour', 'year', 'week', 'date', 'weekdayname','counters_used', 'calendar_year']].dropna() # as Maxbo Nittedal has counter 172.16.164.123 which is not defined
data_sales_date = data_sales_all.groupby(['store_name', 'date', 'type_sales'])['count'].sum().reset_index()
data_sales_date["day_name"] = data_sales_date['date'].dt.weekday_name
data_sales_date = data_sales_date[data_sales_date['day_name'] != "Sunday"] # stores are closed on Sunday's
data_sales_date = pd.pivot_table(data_sales_date, index=['store_name', 'date'], columns='type_sales', fill_value=0).reset_index()
data_sales_date.columns = data_sales_date.columns.droplevel()
data_sales_date.columns = ['store', 'date', 'pos', 'sales']
data_sales_date['all'] = data_sales_date['pos'] + data_sales_date['sales']
data_sales_date = pd.melt(data_sales_date, id_vars=['store', 'date'], value_vars=['pos', 'sales', 'all'], var_name='type_sales', value_name='transactions')
data_sales_date = pd.merge(data_sales_date, sales_group, left_on=['store', 'type_sales'], right_on=['store_name', 'type_sales'], how='right')  # include only necessary type_sales
# conversion rate development
# select only counters for conversion rate index
data_traffic_date = data_traffic[data_traffic['counters_used'] == 'yes'].groupby(['store_name', 'date'])['people_in'].sum().reset_index()
### here should be approximation of the traffic as there are dates with transactions but traffic is zero
conversion_rate = pd.merge(data_sales_date, data_traffic_date, left_on=['store', 'date'], right_on=['store_name', 'date'], how='inner')
# traffic data contains NA's
# explanation: transaction data were send before people counters were installed ['MAXBO SØRUMSAND', 'MAXBO IDEBYGG AS', 'MAXBO MANDAL', 'MAXBO NOTODDEN', 'MAXBO SKANSEN', 'MAXBO SLEMMESTAD', 'MAXBO TRYSIL']
conversion_rate = conversion_rate.replace(np.inf, np.nan).dropna() # as the number transactions and visitors are equal zero we must delete them!
conversion_rate=conversion_rate[conversion_rate['people_in'] != 0]
conversion_rate['week'] = conversion_rate['date'].dt.week
conversion_rate['month'] = conversion_rate['date'].dt.strftime('%B')
conversion_rate = pd.merge(conversion_rate, calendar_weeks, on='date', how='left')
conversion_rate = conversion_rate.groupby(['store', 'week', 'calendar_year', 'month'], as_index = False).agg({"transactions":"sum", "people_in": "sum"}).sort_values(['calendar_year', 'week'])
conversion_rate['conversion_rate'] = round(conversion_rate['transactions']/conversion_rate['people_in'], 4) # contains NANs


today_week=10
today_month='March'
cumulative_weeks="Conversion rate (week 1-"+str(today_week-1)+')'
current_week="Conversion rate (week "+str(today_week)+')'
cumulative_months="Conversion rate (Jan-Feb)"
current_month="Conversion rate (Mar)"

conversion_rate_last_weeks = conversion_rate[conversion_rate['week'] < today_week][['store','conversion_rate']].groupby('store').mean().reset_index().rename(columns={'conversion_rate': cumulative_weeks})
conversion_rate_last_months = conversion_rate[conversion_rate['month'] != today_month].groupby('store')['conversion_rate'].mean().reset_index().rename(columns={'conversion_rate': cumulative_months})
conversion_rate_current_week = conversion_rate[conversion_rate['week'] == today_week][['store','conversion_rate']].rename(columns={'conversion_rate': current_week})
conversion_rate_current_month = conversion_rate[conversion_rate['month'] == today_month][['store','conversion_rate']].rename(columns={'conversion_rate': current_month})
store_development = pd.merge(conversion_rate_current_week, conversion_rate_last_weeks, on='store')
store_development = pd.merge(store_development, conversion_rate_last_months, on='store')
store_development = pd.merge(store_development, conversion_rate_current_month, on='store')
store_development['Conversion rate dev.'] = round((store_development[current_week] - store_development[cumulative_weeks])/store_development[cumulative_weeks], 4)
store_development=store_development.sort_values(by=current_week, ascending=False).rename(columns={'store': 'Store'})

document = Document('input/template.docx')
# define the font family and size for entire file
style = document.styles['Normal']
font = style.font
font.name = 'Avenir Book'
font.size = Pt(10)
# generate stores conversion rate development
table = document.add_table(rows=1, cols=6, style='rm')
col_names=store_development.columns
header_cells = table.rows[0].cells
for k in range(len(col_names)):
    header_cells[k].paragraphs[0].add_run(col_names[k]).bold=True
for i in range(len(store_development)):
    row_cells = table.add_row().cells
    for d in range(len(col_names)):
        if store_development[col_names[d]].dtypes == 'float64':
            row_cells[d].text = str(round(store_development[col_names[d]].values[i]*100, 1))+ '%'
        else:
            row_cells[d].text = str(store_development[col_names[d]].values[i])
document.add_paragraph(text='\n')

# generate stores activity per each region
table = document.add_table(rows=1, cols=7, style='rm')
col_names=data_table_2.columns
header_cells = table.rows[0].cells
for k in range(len(col_names)):
    header_cells[k].paragraphs[0].add_run(col_names[k]).bold=True
for i in range(len(data_table_2)):
    row_cells = table.add_row().cells
    for d in range(len(col_names)):
        if data_table_2[col_names[d]].dtypes == 'float64':
            row_cells[d].text = str(round(data_table_2[col_names[d]].values[i], 2))
        else:
            row_cells[d].text = str(data_table_2[col_names[d]].values[i])
document.add_paragraph(text='\n')

# generate stores activity per each region
for region in sorted(activities['region_name'].unique()):
    activities_region=activities[activities['region_name']==region]
    print('Activity of region %s'%region)
    activities_yes=activities_region[activities_region['is_active']=='yes']
    activities_no=activities_region[activities_region['is_active']=='no']

    document.add_paragraph(text=region)
    table = document.add_table(rows=1, cols=3, style='rm_simple')
    header_cells = table.rows[0].cells
    header_cells[0].text = ''
    header_cells[1].paragraphs[0].add_run('Which stores have been using the system').bold=True
    header_cells[2].paragraphs[0].add_run('Which stores have not been using the system').bold=True
    shade_cells([table.cell(0, 1)], "#00b050")
    shade_cells([table.cell(0, 2)], "#3d393a")
    for i in range(max([len(activities_no), len(activities_yes)])):
        row_cells = table.add_row().cells
        row_cells[0].text=str(i+1)
        if i< len(activities_yes):
            row_cells[1].text=str(activities_yes['store_name'].values[i])
        if i< len(activities_no):
            row_cells[2].text=str(activities_no['store_name'].values[i])
    document.add_paragraph(text='\n')
document.save('output/Maxbo Report.docx')